from django.core.management.base import BaseCommand
from core.models import Course
from docx import Document
from django.utils.text import slugify
import datetime
import os


class Command(BaseCommand):
    args = ''
    help = 'Create a file with the all forum posts and answers for each Course'

    def handle(self, *args, **options):

        courses = Course.objects.filter(status='published')
        if not courses.exists():
            print "There's no published courses to extract"
            return

        dir_name = 'extract_forum'
        if not os.path.exists(dir_name):
            os.makedirs(dir_name)

        print "Generating files for %d courses"
        for course in courses:
            doc = Document()
            doc.add_heading(course.name, level=1)
            doc.add_paragraph()
            for question in course.question_set.filter(hidden=False).order_by('timestamp'):

                user_name = question.user.get_full_name()

                if not user_name:
                    user_name = question.user.username

                p = doc.add_paragraph()
                p.add_run(question.title).bold = True
                p.add_run(' (por ')
                p.add_run(user_name).italic = True
                p.add_run(')')

                p = doc.add_paragraph(question.text)

                for answer in question.answers.order_by('timestamp'):
                    user_name = answer.user.get_full_name()
                    if not user_name:
                        user_name = answer.user.username
                    p = doc.add_paragraph('R: ' + answer.text)
                    p = doc.add_paragraph()
                    p.add_run('por ' + user_name).italic = True

                doc.add_paragraph('_' * 105)
                doc.add_paragraph()

            file_name = os.path.join(dir_name, slugify(course.name) + datetime.datetime.now().strftime('_%Y%m%d%I%M.docx'))

            doc.add_page_break()
            print file_name
            doc.save(file_name)
